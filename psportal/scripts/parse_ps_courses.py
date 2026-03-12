#!/usr/bin/env python3
"""
Parse 'ps courses.docx' to extract courses/levels and upsert to MongoDB.
- Writes to collection 'courses' (PSCourse) and optionally to 'admincourses' (AdminCourse).
- Output: courses-inserted.json

Usage:
  pip install python-docx pymongo python-dotenv
  Export your Google Doc as .docx, save as 'ps courses.docx' in this folder (or set DOCX_PATH).
  Set MONGODB_URI (or MONGO_URI) and DB_NAME=psdb. Then:
  python parse_ps_courses.py

Env: PUSH_ADMIN_COURSES=1 (default) also upserts into admincourses for Admin Dashboard.
"""
import os
import re
import json
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Run: pip install python-docx")
    raise
try:
    from pymongo import MongoClient
    from bson import ObjectId
except ImportError:
    print("Run: pip install pymongo")
    raise
try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = lambda x: None

load_dotenv(Path(__file__).parent.parent / "backend" / ".env")
load_dotenv()

# Config from env
MONGODB_URI = os.getenv("MONGODB_URI") or os.getenv("MONGO_URI") or "mongodb://localhost:27017"
DB_NAME = os.getenv("DB_NAME", "psdb")
COLLECTION = os.getenv("COLLECTION", "courses")
# AdminCourse table (Mongoose model "AdminCourse" -> collection "admincourses")
ADMIN_COLLECTION = os.getenv("ADMIN_COLLECTION", "admincourses")
DOCX_PATH = os.getenv("DOCX_PATH") or str(Path(__file__).parent / "ps courses.docx")
# Set PUSH_ADMIN_COURSES=1 to also upsert into admincourses (default: 1)
PUSH_ADMIN_COURSES = os.getenv("PUSH_ADMIN_COURSES", "1").strip().lower() in ("1", "true", "yes")

# Known subjects (extend as needed)
SUBJECTS = [
    "Analog Electronics", "Digital Electronics", "C Programming", "Data Structure",
    "Database Programming", "Electrical", "Automation", "Communication", "Control Systems",
    "Signals", "Microprocessor", "VLSI", "Embedded", "Power", "Measurements", "Network",
    "Software", "Programming", "Aptitude", "Verbal", "Reasoning",
]

# Patterns: "1 Subject Level - 1A", "2 Subject Level - 1B", "Subject Mock Test 1", "Subject Levels N"
LEVEL_PAT = re.compile(
    r"^(?:\d+[.)]\s*)?(.+?\s+(?:Level\s*-\s*[\w\d]+|Levels?\s+\d+|Mock\s+Test\s+\d+))",
    re.IGNORECASE
)
TOPIC_PAT = re.compile(r"^[\s\-*]*([A-Za-z][^:\n]{5,80})")  # topic-like line
# "0 pts", "300 pts", "Reward points: 300", "300 points"
REWARD_PAT = re.compile(r"(?:reward\s*[:\s]*|points?\s*[:\s]*)?(\d+)\s*(?:pts|points?)?", re.I)
# "Prereq: Level 3", "Prerequisite: Level 1, Level 2", "Prereq: None"
PREREQ_PAT = re.compile(r"prerequisite\s*[:\s]*(.+?)(?:\n|$)|prereq\.?\s*[:\s]*(.+?)(?:\n|$)", re.I | re.DOTALL)
ASSESSMENT_PAT = re.compile(r"assessment\s*(?:type)?\s*[:\s]*([A-Za-z\s]+?)(?:\n|$)", re.I)


def extract_text_from_docx(path: str) -> list[str]:
    if not os.path.isfile(path):
        print(f"File not found: {path}")
        return []
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    lines.append(t)
    return lines


def infer_parent_course(name: str) -> str:
    for s in SUBJECTS:
        if s.lower() in name.lower():
            return s
    if "Level" in name or "Mock" in name:
        parts = name.replace("Level -", " ").replace("Mock Test", " ").split()
        for i, p in enumerate(parts):
            if p and p[0].isalpha() and len(p) > 2:
                return " ".join(parts[:i + 1]) if i > 0 else p
    return name.split(" Level")[0].strip() if " Level" in name else name.split(" Mock")[0].strip()


def is_level_entry(name: str) -> bool:
    return bool(
        re.search(r"Level\s*-\s*[\w\d]+", name, re.I)
        or re.search(r"Mock\s+Test\s+\d+", name, re.I)
        or re.search(r"Levels?\s+\d+", name, re.I)
    )


def build_description_from_next_lines(lines: list[str], start: int, max_lines: int = 2) -> str:
    topics = []
    for i in range(start + 1, min(start + 1 + max_lines, len(lines))):
        line = lines[i]
        if re.match(r"^\d+[.)]\s*", line) or LEVEL_PAT.match(line):
            break
        if len(line) > 10 and not line.startswith("Attempts") and "reward" not in line.lower():
            topics.append(line[:100].strip())
    return "; ".join(topics)[:200] if topics else ""


def extract_metadata_from_next_lines(
    lines: list[str], start: int, max_lines: int = 12,
    level_name: str = "", parent_course: str = ""
) -> dict:
    """From lines following a level heading, extract rewardPoints, assessmentType, prereqLevelNames (from doc)."""
    reward_points = 0
    assessment_type = ""
    prereq_raw = ""
    seen_explicit_assessment = False
    for i in range(start + 1, min(start + 1 + max_lines, len(lines))):
        line = lines[i]
        if re.match(r"^\d+[.)]\s*", line) or LEVEL_PAT.match(line):
            break
        line_stripped = line.strip()
        line_lower = line_stripped.lower()
        # Reward: "0 pts", "300 pts", "Reward points: 300", or standalone number
        if re.search(r"\d+\s*pts?\b", line_lower) or "reward" in line_lower or ("point" in line_lower and "attempt" not in line_lower):
            m = REWARD_PAT.search(line)
            if m:
                reward_points = int(m.group(1))
        elif reward_points == 0 and line_stripped.isdigit() and len(line_stripped) <= 5:
            reward_points = int(line_stripped)
        # Prerequisite: "Prereq: Level 3", "Prerequisite: Level 1, Level 2", "None (first level)"
        if "prerequisite" in line_lower or "prereq" in line_lower:
            m = PREREQ_PAT.search(line)
            if m:
                raw = (m.group(1) or m.group(2) or "").strip()
                if raw and "none" not in raw.lower():
                    prereq_raw = raw
            elif "none" in line_lower or "first level" in line_lower:
                prereq_raw = "none"
            else:
                prereq_raw = line_stripped
        elif "none (first" in line_lower and not prereq_raw:
            prereq_raw = "none"
        # Assessment type: use doc value; if course/level name contains "Programming", prefer Programming over MCQ
        if "assessment" in line_lower:
            m = ASSESSMENT_PAT.search(line)
            if m:
                assessment_type = (m.group(1) or "").strip()
                seen_explicit_assessment = True
            else:
                rest = re.sub(r"assessment\s*(?:type)?\s*[:\s]*", "", line_lower, flags=re.I).strip()
                if rest and len(rest) < 30:
                    assessment_type = rest.title()
                    seen_explicit_assessment = True
        elif len(line_stripped) < 40:
            if line_lower == "programming":
                assessment_type = "Programming"
                seen_explicit_assessment = True
            elif line_lower == "mcq":
                assessment_type = "MCQ"
                seen_explicit_assessment = True
            elif not assessment_type:
                for at in ("Programming", "MCQ", "Manual Grading", "Coding"):
                    if at.lower() == line_lower or (len(line_lower) > 2 and at.lower() in line_lower):
                        assessment_type = at
                        seen_explicit_assessment = True
                        break
    # Parse prereq into list: "Level 3" -> ["Level 3"], "Level 1, Level 2" -> ["Level 1", "Level 2"]
    prereq_level_names = []
    if prereq_raw and "none" not in prereq_raw.lower():
        for part in re.split(r"[,;]|\band\b", prereq_raw, flags=re.I):
            p = part.strip()
            if p and len(p) > 1 and not p.lower().startswith("("):
                prereq_level_names.append(p)
    # Default assessment: if course/subject is "Programming", use "Programming" unless doc said something else
    combined_name = (level_name + " " + parent_course).lower()
    if not assessment_type or (not seen_explicit_assessment and "programming" in combined_name):
        assessment_type = "Programming" if "programming" in combined_name else (assessment_type or "MCQ")
    return {
        "rewardPoints": reward_points,
        "assessmentType": (assessment_type or "MCQ").strip() or "MCQ",
        "prereqLevelNames": prereq_level_names,
    }


def parse_courses(lines: list[str]) -> list[dict]:
    entries = []
    seen = set()
    i = 0
    while i < len(lines):
        line = lines[i]
        m = LEVEL_PAT.match(line)
        if m:
            name = m.group(1).strip()
            name = re.sub(r"^\d+[.)]\s*", "", name).strip()
            if not name or name in seen:
                i += 1
                continue
            parent = infer_parent_course(name)
            seen.add(name)
            desc = build_description_from_next_lines(lines, i)
            meta = extract_metadata_from_next_lines(lines, i, 12, level_name=name, parent_course=parent)
            level = is_level_entry(name)
            entries.append({
                "name": name,
                "description": desc or f"Course: {name}",
                "status": "Active",
                "level": level,
                "parentCourse": parent,
                "prereq": [],
                "rewardPoints": meta["rewardPoints"],
                "assessmentType": meta["assessmentType"],
                "prereqLevelNames": meta["prereqLevelNames"],
            })
        else:
            for subj in SUBJECTS:
                if subj.lower() in line.lower() and "Level" in line:
                    name = line.strip()
                    name = re.sub(r"^\d+[.)]\s*", "", name).strip()
                    if name and name not in seen and len(name) > 5:
                        seen.add(name)
                        meta = extract_metadata_from_next_lines(lines, i, 12, level_name=name, parent_course=subj)
                        entries.append({
                            "name": name,
                            "description": "",
                            "status": "Active",
                            "level": True,
                            "parentCourse": subj,
                            "prereq": [],
                            "rewardPoints": meta["rewardPoints"],
                            "assessmentType": meta["assessmentType"],
                            "prereqLevelNames": meta["prereqLevelNames"],
                        })
                    break
        i += 1
    return entries


def group_entries_by_subject(entries: list[dict]) -> dict:
    """Group level entries by parent course (subject). One course per subject with levels inside."""
    grouped = {}
    for e in entries:
        parent = (e.get("parentCourse") or "").strip()
        if not parent:
            continue
        if parent not in grouped:
            grouped[parent] = []
        grouped[parent].append(e)
    # Sort levels by name within each subject
    for parent in grouped:
        grouped[parent].sort(key=lambda x: (x.get("name") or ""))
    return grouped


def _topics_from_description(desc: str, max_items: int = 15) -> list[str]:
    """Parse topics from description: newline or semicolon separated."""
    if not (desc or "").strip():
        return []
    raw = (desc or "").replace(";", "\n").split("\n")
    topics = [t.strip() for t in raw if t.strip()][:max_items]
    return topics


def _resolve_prereq_indices(level_entries: list[dict], prereq_names: list[str], current_index: int) -> list[int]:
    """Map prereq level names (from doc) to indices in level_entries. E.g. 'Level 3' -> index of '... Level - 3'."""
    if not prereq_names:
        return []
    indices = []
    names_lower = [((e.get("name") or "").lower()) for e in level_entries]
    for pr in prereq_names:
        pr_lower = (pr or "").strip().lower()
        if not pr_lower or pr_lower == "none":
            continue
        pr_norm = re.sub(r"\s*-\s*", " ", pr_lower)
        for j, lev_name in enumerate(names_lower):
            if j >= current_index:
                break
            lev_norm = re.sub(r"\s*-\s*", " ", lev_name)
            if pr_lower in lev_name or pr_lower in lev_norm or pr_norm in lev_norm or lev_name.endswith(pr_lower):
                indices.append(j)
                break
            # "Level 3" vs "Database Programming Level - 3" -> extract number 3
            pr_num = re.search(r"level\s*[-]?\s*(\d+[a-z]?)", pr_lower)
            lev_num = re.search(r"level\s*[-]?\s*(\d+[a-z]?)", lev_name)
            if pr_num and lev_num and pr_num.group(1) == lev_num.group(1):
                indices.append(j)
                break
    return sorted(set(indices))


def _get_prereq_indices_for_level(level_entries: list[dict], entry: dict, idx: int) -> list[int]:
    """Prereq indices: from doc (prereqLevelNames) if present, else infer previous level required."""
    prereq_names = entry.get("prereqLevelNames") or []
    if prereq_names:
        return _resolve_prereq_indices(level_entries, prereq_names, idx)
    if idx <= 0:
        return []
    return [idx - 1]


def _level_number_from_name(level_name: str) -> int:
    """Parse level number from name: 'Level - 1', 'Level - 2', 'Level - 3A', 'Level 1' -> 1, 2, 3. Default 1."""
    if not level_name:
        return 1
    m = re.search(r"level\s*[-]?\s*(\d+)", (level_name or "").lower())
    return int(m.group(1)) if m else 1


def _default_reward_points(assessment_type: str, level_index: int, doc_value: int, level_name: str = "") -> int:
    """Programming: Level 1=300, Level 2=600, Level 3+=900. MCQ: all 100. Use level number from name, not index."""
    if doc_value > 0:
        return doc_value
    at = (assessment_type or "").strip().lower()
    if at == "programming":
        num = _level_number_from_name(level_name) if level_name else (level_index + 1)
        if num == 1:
            return 300
        if num == 2:
            return 600
        return 900  # Level 3 onwards
    return 100  # MCQ and others: all levels 100


def build_level_objects_for_admin(level_entries: list[dict]) -> list[dict]:
    """AdminCourse levels: name, rewardPoints (doc or default by Programming/MCQ rule), prereq, assessmentType, topics."""
    out = []
    for idx, e in enumerate(level_entries):
        desc = (e.get("description") or "").strip()
        topics = _topics_from_description(desc)
        prereq_indices = _get_prereq_indices_for_level(level_entries, e, idx)
        assessment_type = (e.get("assessmentType") or "MCQ").strip() or "MCQ"
        doc_pts = int(e.get("rewardPoints") or 0)
        level_name = e.get("name", "")
        reward_pts = _default_reward_points(assessment_type, idx, doc_pts, level_name)
        out.append({
            "name": level_name,
            "rewardPoints": reward_pts,
            "prerequisiteLevelIndex": prereq_indices[0] if prereq_indices else -1,
            "prerequisiteLevelIndices": prereq_indices,
            "assessmentType": assessment_type,
            "topics": topics,
        })
    return out


def build_level_objects_for_ps(level_entries: list[dict]) -> list[dict]:
    """PSCourse levels: same shape; reward points by Programming/MCQ rule when doc has 0."""
    out = []
    for idx, e in enumerate(level_entries):
        desc = (e.get("description") or "").strip()
        topics = _topics_from_description(desc)
        prereq_indices = _get_prereq_indices_for_level(level_entries, e, idx)
        assessment_type = (e.get("assessmentType") or "MCQ").strip() or "MCQ"
        doc_pts = int(e.get("rewardPoints") or 0)
        level_name = e.get("name", "")
        reward_pts = _default_reward_points(assessment_type, idx, doc_pts, level_name)
        out.append({
            "name": level_name,
            "description": desc[:500] if desc else "",
            "rewardPoints": reward_pts,
            "prerequisiteLevelIndex": prereq_indices[0] if prereq_indices else -1,
            "prerequisiteLevelIndices": prereq_indices,
            "assessmentType": assessment_type,
            "topics": topics,
        })
    return out


def to_mongo_doc(entry: dict) -> dict:
    """Legacy flat doc (for JSON export only)."""
    now = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
    return {
        "_id": str(ObjectId()),
        "name": entry["name"],
        "description": entry.get("description", ""),
        "status": entry.get("status", "Active"),
        "level": entry.get("level", False),
        "parentCourse": entry.get("parentCourse", ""),
        "prereq": entry.get("prereq") or [],
        "createdAt": now,
        "updatedAt": now,
        "__v": 0,
    }


def build_course_docs_with_levels(grouped: dict) -> tuple[list[dict], list[dict]]:
    """Returns (ps_course_docs, admin_course_docs). One doc per subject with levels array."""
    ps_docs = []
    admin_docs = []
    for subject_name, level_entries in grouped.items():
        if not level_entries:
            continue
        first_desc = (level_entries[0].get("description") or "").strip() or f"Course: {subject_name}"
        admin_docs.append({
            "name": subject_name,
            "description": first_desc,
            "status": "Active",
            "type": "PS",
            "course_logo": "",
            "level": "",
            "activity_points": 0,
            "reward_points": 0,
            "faculty": "",
            "prerequisites": [],
            "levels": build_level_objects_for_admin(level_entries),
        })
        ps_docs.append({
            "name": subject_name,
            "description": first_desc,
            "status": "Active",
            "level": True,
            "parentCourse": "",
            "prereq": [],
            "levels": build_level_objects_for_ps(level_entries),
        })
    return ps_docs, admin_docs


def upsert_courses_with_levels(ps_docs: list[dict], admin_docs: list[dict], old_level_names: set[str]) -> tuple[int, int]:
    """Upsert one course per subject with levels array. Remove old flat level documents."""
    client = MongoClient(MONGODB_URI)
    db = client[DB_NAME]
    now_iso = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
    now = datetime.now(timezone.utc)
    subject_names = set(d["name"] for d in ps_docs)

    # Remove old flat entries (e.g. "C Programming Level - 1A") so we don't keep 244 separate docs
    to_remove = old_level_names - subject_names
    if to_remove:
        db[COLLECTION].delete_many({"name": {"$in": list(to_remove)}})
        if admin_docs:
            db[ADMIN_COLLECTION].delete_many({"name": {"$in": list(to_remove)}})

    # Upsert PSCourse (courses)
    n_ps = 0
    for d in ps_docs:
        set_fields = {
            "description": d["description"],
            "status": d["status"],
            "level": d["level"],
            "parentCourse": d["parentCourse"],
            "prereq": d["prereq"],
            "levels": d["levels"],
            "updatedAt": now,
        }
        db[COLLECTION].update_one(
            {"name": d["name"]},
            {"$set": set_fields, "$setOnInsert": {"_id": ObjectId(), "name": d["name"], "createdAt": now, "__v": 0}},
            upsert=True,
        )
        n_ps += 1

    # Upsert AdminCourse (admincourses)
    n_admin = 0
    for d in admin_docs:
        set_fields = {
            "description": d["description"],
            "status": d["status"],
            "type": d["type"],
            "course_logo": d["course_logo"],
            "level": d["level"],
            "activity_points": d["activity_points"],
            "reward_points": d["reward_points"],
            "faculty": d["faculty"],
            "prerequisites": d["prerequisites"],
            "levels": d["levels"],
            "updatedAt": now,
        }
        db[ADMIN_COLLECTION].update_one(
            {"name": d["name"]},
            {"$set": set_fields, "$setOnInsert": {"name": d["name"], "createdAt": now}},
            upsert=True,
        )
        n_admin += 1

    client.close()
    return n_ps, n_admin


def main():
    print(f"Reading: {DOCX_PATH}")
    lines = extract_text_from_docx(DOCX_PATH)
    if not lines:
        print("No content. Create sample entries.")
        entries = [
            {"name": "Analog Electronics Level - 1A", "description": "Selection of Current Limiting Resistors", "status": "Active", "level": True, "parentCourse": "Analog Electronics", "prereq": []},
        ]
    else:
        entries = parse_courses(lines)
        print(f"Parsed {len(entries)} level entries from docx.")

    if not entries:
        print("No entries to insert.")
        return

    grouped = group_entries_by_subject(entries)
    print(f"Grouped into {len(grouped)} courses (subjects) with levels inside.")
    ps_docs, admin_docs = build_course_docs_with_levels(grouped)

    out_path = Path(__file__).parent / "courses-inserted.json"
    export = []
    for d in ps_docs:
        # Every level in JSON has full fields: name, description, rewardPoints, prerequisiteLevelIndex, prerequisiteLevelIndices, assessmentType, topics
        levels_export = []
        for lev in d.get("levels") or []:
            levels_export.append({
                "name": lev.get("name", ""),
                "description": lev.get("description", ""),
                "rewardPoints": int(lev.get("rewardPoints") or 0),
                "prerequisiteLevelIndex": int(lev.get("prerequisiteLevelIndex") if lev.get("prerequisiteLevelIndex") is not None else -1),
                "prerequisiteLevelIndices": list(lev.get("prerequisiteLevelIndices") or []),
                "assessmentType": (lev.get("assessmentType") or "MCQ").strip() or "MCQ",
                "topics": list(lev.get("topics") or []),
            })
        export.append({
            "name": d["name"],
            "description": d["description"],
            "status": d["status"],
            "level": d["level"],
            "parentCourse": d["parentCourse"],
            "prereq": d["prereq"],
            "levels": levels_export,
            "levelsCount": len(levels_export),
        })
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(export, f, indent=2, ensure_ascii=False)
    print(f"Wrote {len(export)} courses (with levels) to {out_path}")

    if MONGODB_URI and ("localhost" in MONGODB_URI or "mongodb" in MONGODB_URI):
        try:
            old_level_names = {e["name"] for e in entries}
            n_ps, n_admin = upsert_courses_with_levels(
                ps_docs,
                admin_docs if PUSH_ADMIN_COURSES else [],
                old_level_names,
            )
            if PUSH_ADMIN_COURSES:
                print(f"MongoDB: {n_ps} courses (with levels) in {DB_NAME}.{COLLECTION}, {n_admin} in {DB_NAME}.{ADMIN_COLLECTION}")
            else:
                print(f"MongoDB: {n_ps} courses (with levels) in {DB_NAME}.{COLLECTION}")
        except Exception as err:
            print(f"MongoDB upsert failed: {err}")
    else:
        print("Set MONGODB_URI/MONGO_URI to upsert to MongoDB.")


if __name__ == "__main__":
    main()
